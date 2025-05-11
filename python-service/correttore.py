"""
Corregge grammatica e ortografia in **qualsiasi** parte di un documento word (.docx) e produce un report Markdown con tutte le modifiche.

* testo normale, anche in tabelle nidificate
* header & footer
* testo delle note a piè di pagina (footnotes)
* caselle di testo / forme (<w:txbxContent>)

Mantiene corsivi, grassetti, sottolineature, riferimenti di nota, ecc.
Richiede **python-docx ≥ 0.8.11**.
"""

from __future__ import annotations

import os
import re
import time
import json
import shutil
import collections
import asyncio
import zipfile
import tiktoken
import tempfile
from collections import defaultdict
from common import tokenize, WORD_RE
from pathlib import Path
from dotenv import load_dotenv
from docx import Document
from docx.oxml.ns import qn
from dataclasses import dataclass
from difflib import SequenceMatcher
from docx.text.paragraph import Paragraph
from openai import OpenAI
from openai import AsyncOpenAI
from typing import Dict, Iterable, List, Optional, Tuple
from reports import write_markdown_report, write_glossary_report
from utils_openai import _OPENAI_MODEL as OPENAI_MODEL
from utils_openai import get_corrections_async, get_corrections_sync

# ───────────────────────── CONFIGURAZIONE ────────────────────────────
load_dotenv()
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
if not OPENAI_API_KEY:
    raise ValueError("API key di OpenAI non trovata. Imposta OPENAI_API_KEY nel tuo ambiente.")
# Lunghezza massima di contesto (in token) accettata in un singolo prompt
MAX_TOKENS_GPT4O_MINI = 10000

try:
    ENC = tiktoken.encoding_for_model(OPENAI_MODEL)
except KeyError:
    # fallback universale, compatibile con GPT-4/3.5
    ENC = tiktoken.get_encoding("cl100k_base")
# ─────────────────────────────────────────────────────────────────────

NAME_RE = re.compile(r"\b(?:[A-Z][a-z]{2,}|[A-Z]{2,})\w*\b")

# ╭─────────────────────────── Note a piè di pagina ▾──────────────────────────╮
from copy import deepcopy
import os
from lxml import etree

# nuova importazione della utility condivisa


def correggi_footnotes_xml(docx_path: Path,
                           client,
                           glossary: set[str] | None = None) -> None:
    """
    Corregge le note a piè di pagina contenute in word/footnotes.xml
    (refusi, ortografia, punteggiatura) preservando la formattazione
    run-per-run del documento Word.
    """
    glossary = glossary or set()           # se None, usa set vuoto
    tmp_dir = Path(tempfile.mkdtemp(prefix="docx_", dir=tempfile.gettempdir()))

    # 1) Estrai il .docx in una cartella temporanea --------------------
    with zipfile.ZipFile(docx_path, "r") as zf:
        zf.extractall(tmp_dir)

    footnotes_file = os.path.join(tmp_dir, "word", "footnotes.xml")
    if not os.path.exists(footnotes_file):
        shutil.rmtree(tmp_dir)
        return                              # documento senza note

    ns   = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    tree = etree.parse(footnotes_file)

    # 2) Scorri ogni <w:footnote> (escludendo i separator) -------------
    for foot in tree.xpath("//w:footnote[not(@w:type='separator')]", namespaces=ns):
        txt_nodes  = foot.xpath(".//w:t", namespaces=ns)
        full_text  = "".join(n.text or "" for n in txt_nodes)

        if not full_text.strip():
            continue                       # nota vuota → salta

        # 2a) Chiamata OpenAI (una sola riga, via utility) -------------
        corr = get_corrections_sync(
            payload_json = json.dumps([{"id": 0, "txt": full_text}], ensure_ascii=False),
            client       = client,
            glossary     = glossary,
            context      = "",
        )
        corrected = corr[0]["txt"]

        if corrected == full_text:
            continue                       # nessuna correzione

        # 2b) Redistribuisci token corretti nei singoli <w:t> ----------
        orig_tok   = tokenize(full_text)
        corr_tok   = tokenize(corrected)
        mapping    = align_tokens(orig_tok, corr_tok)
        starts_orig = token_starts(orig_tok)

        # mappa carattere→indice nodo
        char2node = []
        for idx, n in enumerate(txt_nodes):
            char2node.extend([idx] * len(n.text or ""))

        tok_per_node = defaultdict(list)
        for ref_idx, tok in mapping:
            if ref_idx is None or ref_idx >= len(starts_orig):
                node_idx = 0
            else:
                char_pos = starts_orig[ref_idx]
                node_idx = char2node[min(char_pos, len(char2node) - 1)]
            tok_per_node[node_idx].append(tok)

        # scrivi il testo nei nodi preservando la partizione originale
        for idx, n in enumerate(txt_nodes):
            n.text = "".join(tok_per_node.get(idx, []))

        # 2c) Se si perde la formattazione del primo run, ristabiliscila
        first_with_txt = next((n for n in txt_nodes if n.text and n.text.strip()), None)
        if first_with_txt is not None:
            run_first       = first_with_txt.getparent()
            has_rPr_first   = run_first.find("./w:rPr", namespaces=ns)
            has_bold_first  = run_first.xpath("./w:rPr/w:b", namespaces=ns)

            if not has_bold_first:
                for n2 in txt_nodes:
                    if n2 is first_with_txt or not (n2.text and n2.text.strip()):
                        continue
                    rPr_other = n2.getparent().find("./w:rPr", namespaces=ns)
                    if rPr_other is not None and list(rPr_other):
                        if has_rPr_first is None:
                            has_rPr_first = etree.SubElement(run_first, qn("w:rPr"))
                        for child in rPr_other:
                            has_rPr_first.append(deepcopy(child))
                        break   # una sola copia è sufficiente

    # 3) Salva il nuovo footnotes.xml e ricompatta il .docx -------------
    tree.write(footnotes_file, xml_declaration=True, encoding="utf-8")

    tmp_docx = docx_path.with_suffix(".tmp")
    with zipfile.ZipFile(tmp_docx, "w") as zf:
        for root, _, files in os.walk(tmp_dir):
            for f in files:
                fullpath = os.path.join(root, f)
                arcname  = os.path.relpath(fullpath, tmp_dir)
                zf.write(fullpath, arcname)

    shutil.move(tmp_docx, docx_path)        # sovrascrive l’originale
    shutil.rmtree(tmp_dir)
    print("✏️  Note a piè di pagina corrette e formattazione preservata")

# ╭─────────────────────────── Utility token ▾──────────────────────────╮
def tokenize(text: str) -> List[str]:
    """Tokenizzazione grezza (regex) — sufficiente per stimare la
    lunghezza in token senza dipendenze esterne."""
    return WORD_RE.findall(text)

def count_tokens(text: str) -> int:
    """Conta i token reali secondo l’encoding del modello."""
    return len(ENC.encode(text or ""))

def token_starts(tokens: List[str]) -> List[int]:
    pos = 0
    starts = []
    for tok in tokens:
        starts.append(pos)
        pos += len(tok)
    return starts
# ╰──────────────────────────────────────────────────────────────────────╯

# ╭─────────────────────────── Data-model modifiche ▾────────────────────╮
@dataclass
class Modification:
    par_id: int
    original: str
    corrected: str
# ╰──────────────────────────────────────────────────────────────────────╯

# ╭─────────────────────────── Chunking ▾────────────────────────────────╮
def chunk_paragraph_objects(
    paragraphs: List[Paragraph],
    max_tokens: int = MAX_TOKENS_GPT4O_MINI,
) -> List[List[Paragraph]]:
    """Dividi la lista di oggetti Paragraph in blocchi < max_tokens."""
    chunks: List[List[Paragraph]] = []
    current: List[Paragraph] = []
    current_tokens = 0

    for p in paragraphs:
        para_tokens = count_tokens(p.text)

        if current and current_tokens + para_tokens > max_tokens:
            chunks.append(current)
            current = [p]
            current_tokens = para_tokens
        else:
            current.append(p)
            current_tokens += para_tokens

        # paragrafo singolo > soglia
        if not current and para_tokens > max_tokens:
            chunks.append([p])
            current_tokens = 0

    if current:
        chunks.append(current)
    return chunks
# ╰──────────────────────────────────────────────────────────────────────╯

# ╭─────────────────────────── Diff token-level ▾────────────────────────╮
def align_tokens(orig: List[str], corr: List[str]) -> List[Tuple[Optional[int], str]]:
    sm = SequenceMatcher(a=orig, b=corr, autojunk=False)
    out: List[Tuple[Optional[int], str]] = []
    for op, i1, i2, j1, j2 in sm.get_opcodes():
        if op == "equal":
            for k in range(i2 - i1):
                out.append((i1 + k, corr[j1 + k]))
        elif op == "replace":
            inherit = i1 if i1 < len(orig) else None
            for k in range(j2 - j1):
                out.append((inherit, corr[j1 + k]))
        elif op == "insert":
            inherit = i1 - 1 if i1 > 0 else None
            for k in range(j1, j2):
                out.append((inherit, corr[k]))
    return out
# ╰──────────────────────────────────────────────────────────────────────╯

# ╭─────────────────────────── Helpers copy RUN ▾────────────────────────╮
def copy_rPr(src_rPr, dest_run):
    if src_rPr is None:
        return
    dest = dest_run._r.get_or_add_rPr()
    dest.clear()
    for child in src_rPr:
        dest.append(deepcopy(child))


def clone_run(src_run, paragraph):
    new_run = paragraph.add_run("")
    copy_rPr(src_run._r.rPr, new_run)
    for child in src_run._r:
        if child.tag != qn("w:t"):
            new_run._r.append(deepcopy(child))
    return new_run
# ╰──────────────────────────────────────────────────────────────────────╯

# ─────────────────────────── NUOVO BLOCCO paragrafi multipli ──────────────────────────────
def apply_correction_to_paragraph(
    p: Paragraph,
    corrected: str,
    mods: List[Modification],
    par_id: int,
    glossary: set[str],
):
    """
    Sovrascrive il paragrafo `p` con il testo già corretto
    preservandone la formattazione run-per-run.
    """
    original = p.text
    if corrected == original:
        return

    mods.append(Modification(par_id, original, corrected))

    # === ricostruzione dei run (stessa logica di prima) ===============
    orig_tok  = tokenize(original)
    corr_tok  = tokenize(corrected)
    mapping   = align_tokens(orig_tok, corr_tok)
    starts    = token_starts(orig_tok)
    char_run  = char_to_run_map(p)

    tokens_per_run: Dict[int, List[str]] = defaultdict(list)
    last_run_idx: Optional[int] = None
    for ref_idx, tok in mapping:
        if not char_run:
            run_idx = 0
        elif ref_idx is None:
            run_idx = last_run_idx if last_run_idx is not None else char_run[0]
        else:
            pos     = starts[ref_idx]
            run_idx = char_run[pos] if pos < len(char_run) else char_run[-1]
        last_run_idx = run_idx
        tokens_per_run[run_idx].append(tok)

    old_runs = list(p.runs)
    p._p.clear_content()

    for idx, run in enumerate(old_runs):
        toks = tokens_per_run.get(idx, [])
        if run.text:
            if not toks:
                continue
            new_run = clone_run(run, p)
            # 💥 work-around: se add_run ha restituito None rigenera il run
            if new_run is None:
                new_run = p.add_run("")
            new_run.text = "".join(toks)
        else:
            p._p.append(deepcopy(run._r))
                # --- aggiorna dinamicamente il glossario -------------------------
    for name in NAME_RE.findall(corrected):
        glossary.add(name)


# ╭─────────────────── Funzione generica di correzione ─────────────────╮
async def correct_paragraph_group(
    paragraphs:   list[Paragraph],
    all_paras:    list[Paragraph],
    start_par_id: int,
    client:       AsyncOpenAI,
    glossary:     set[str],
    mods:         list[Modification],
    context_size: int = 3,          # quanti paragrafi di contesto ricavare
):
    """
    Corregge un gruppo di Paragraph mantenendo formattazione e glossario.

    · paragraphs   : la “sezione” da correggere (chunk, note, header…)
    · all_paras    : lista completa per calcolare il contesto
    · start_par_id : id del primo paragrafo nel documento (1-based)
    · client       : istanza AsyncOpenAI condivisa
    · glossary     : set globale dei nomi canonici
    · mods         : lista in cui accumulare le modifiche per il report
    """

    # 1.  CONTEXTO – ultimi `context_size` paragrafi prima di questo blocco
    ctx_start = max(0, start_par_id - context_size - 1)
    context = "\n".join(p.text for p in all_paras[ctx_start : start_par_id - 1])

    # 2.  PAYLOAD JSON
    payload = [{"id": i, "txt": p.text} for i, p in enumerate(paragraphs)]
    payload_json = json.dumps(payload, ensure_ascii=False)

    # 3.  MESSAGGI
    messages = build_messages(context, payload_json, glossary)

    # 4. CHIAMATA OpenAI (ora delegata alla utility)
    corr_list = await get_corrections_async(
        payload_json = payload_json,   # JSON già costruito al punto 2
        client       = client,         # l’istanza AsyncOpenAI passata alla funzione
        glossary     = glossary,       # il set di nomi canonici
        context      = context,        # le righe di contesto calcolate al punto 1
    )

    # 5.  APPLICA LE CORREZIONI
    corr_by_id = {d["id"]: d["txt"] for d in corr_list}

    for local_id, p in enumerate(paragraphs):
        apply_correction_to_paragraph(
            p,
            corr_by_id.get(local_id, p.text),
            mods,
            start_par_id + local_id,
            glossary,                 # 👈 nuovo argomento
        )
    # 5-bis  verifica che il modello non abbia “accorciato” troppo
    for local_id, p in enumerate(paragraphs):
        orig_tok = tokenize(p.text)
        corr_tok = tokenize(corr_by_id.get(local_id, p.text))

        # se ha eliminato >2 % dei token, rigenera il chunk con GPT-4o «full»
        if len(orig_tok) and (len(orig_tok) - len(corr_tok)) / len(orig_tok) > 0.02:
            # puoi fare un retry con il modello maggiore o lanciare errore
            raise RuntimeError(
                f"Chunk {start_par_id+local_id}: eliminazione sospetta "
                f"({len(orig_tok)-len(corr_tok)} token)."
            )

# ╰──────────────────────────────────────────────────────────────────────╯
# ╭──────────────────── Wrapper: corpo principale ─────────────────────╮
async def fix_body_chunks(
    async_client: AsyncOpenAI,
    all_paras:    list[Paragraph],
    para_chunks:  list[list[Paragraph]],
    start_id:     int,
    mods:         list[Modification],
    glossary:     set[str],
):
    """
    Scorre tutti i chunk creati da `chunk_paragraph_objects` e li manda
    in parallelo a `correct_paragraph_group`.
    """
    tasks = []
    par_id = start_id
    for chunk in para_chunks:
        tasks.append(
            correct_paragraph_group(
                paragraphs   = chunk,
                all_paras    = all_paras,
                start_par_id = par_id,
                client       = async_client,
                glossary     = glossary,
                mods         = mods,
            )
        )
        par_id += len(chunk)

    await asyncio.gather(*tasks)
# ╰──────────────────────────────────────────────────────────────────────╯

# ╭─────────────────────────── Map char→run ▾────────────────────────────╮
def char_to_run_map(paragraph) -> List[int]:
    mapping: List[int] = []
    for idx, run in enumerate(paragraph.runs):
        mapping.extend([idx] * len(run.text))
    return mapping
# ╰──────────────────────────────────────────────────────────────────────╯

# ───────────────────────── traversal del documento ─────────────────────────────
def iter_body_paragraphs(container) -> Iterable[Paragraph]:
    for para in container.paragraphs:
        yield para
    for tbl in getattr(container, "tables", []):
        for row in tbl.rows:
            for cell in row.cells:
                yield from iter_body_paragraphs(cell)


def iter_footnote_paragraphs(doc: Document) -> Iterable[Paragraph]:
    fpart = getattr(doc.part, "footnotes_part", None)
    if fpart:
        for footnote in fpart.footnotes:
            for para in footnote.paragraphs:
                yield para


def iter_header_footer_paragraphs(doc: Document) -> Iterable[Paragraph]:
    for sect in doc.sections:
        for hf in (sect.header, sect.footer):
            if hf:
                yield from iter_body_paragraphs(hf)


def iter_textbox_paragraphs(doc: Document) -> Iterable[Paragraph]:
    parts = [doc.part]
    for sect in doc.sections:
        parts.extend([sect.header.part, sect.footer.part])
    fpart = getattr(doc.part, "footnotes_part", None)
    if fpart:
        parts.append(fpart)
    for part in parts:
        root = part._element
        for txbx in root.xpath('.//*[local-name()="txbxContent"]'):
            for p_el in txbx.xpath('.//*[local-name()="p"]'):
                yield Paragraph(p_el, part)


def iter_all_paragraphs(doc: Document) -> Iterable[Paragraph]:
    yield from iter_body_paragraphs(doc)
    yield from iter_footnote_paragraphs(doc)
    yield from iter_header_footer_paragraphs(doc)
    yield from iter_textbox_paragraphs(doc)
# ──────────────────────────────────────────────────────────────────────

# ───────────────────────── entry-point con logging ────────────────────
def process_doc(inp: Path, out: Path):
    """Versione *sincrona* per la CLI; delega alla async con asyncio.run."""
    asyncio.run(process_doc_async(inp, out))
    
async def process_doc_async(inp: Path, out: Path):
    """Identica a process_doc ma *senza* asyncio.run: la useremo da FastAPI."""
    doc = Document(inp)

    # Tutti i paragrafi
    all_paras = list(iter_all_paragraphs(doc))

    # Glossario iniziale
    global GLOSSARY
    name_counts = collections.Counter(
        n for p in all_paras for n in NAME_RE.findall(p.text)
    )
    GLOSSARY = {w for w, c in name_counts.items() if c >= 2}

    # Chunk
    para_chunks = chunk_paragraph_objects(all_paras, max_tokens=300)
    print(f"🔍  Rilevati {len(para_chunks)} chunk (limite {MAX_TOKENS_GPT4O_MINI} token).")

    mods: list[Modification] = []

    # --- corpo documento (async) ---------------------------------------
    async_client = AsyncOpenAI(api_key=OPENAI_API_KEY)
    await fix_body_chunks(
        async_client, all_paras, para_chunks, 1, mods, GLOSSARY
    )

    # --- salvataggi sincroni ------------------------------------------
    doc.save(out)
    client_sync = OpenAI(api_key=OPENAI_API_KEY)
    correggi_footnotes_xml(out, client_sync)
    write_markdown_report(mods, out)
    write_glossary_report(GLOSSARY, all_paras, out)

# ───────────────────────────────────────────────────────────────────────

# ╭────────────────────────── Prompt & builder messaggi ─────────────────────────╮
SYSTEM_MSG_BASE = """
Sei un correttore di bozze madrelingua italiano con decenni di esperienza.

• Correggi **solo** refusi, errori ortografici / grammaticali e punteggiatura.  
• Non eliminare, spostare o accorciare parole, frasi o capoversi.  
• Non riformulare lo stile; se una parte è già corretta, lasciala invariata.

NOMI / TERMINI FANTASY ↓  
Se trovi varianti ortografiche dei nomi presenti nell’elenco seguente,
uniforma la grafia a quella esatta dell’elenco.

OUTPUT: restituisci **SOLO JSON** con la chiave `'corr'`
( lista di {id:int, txt:str} ) — niente testo extra.
"""

def build_messages(context: str, payload_json: str, glossary: set[str]) -> list[dict]:
    """
    Crea i tre messaggi da mandare a OpenAI:
        1. system    → vincoli + lista dei nomi “canonici”
        2. assistant → contesto di righe precedenti (NON va modificato)
        3. user      → JSON dei paragrafi da correggere
    """
    system_msg = SYSTEM_MSG_BASE + "\nLista: " + ", ".join(sorted(glossary))

    return [
        {"role": "system",    "content": system_msg},
        {"role": "assistant", "content": "Contesto (NON modificare):\n" + context},
        {"role": "user",      "content": payload_json},
    ]
# ╰──────────────────────────────────────────────────────────────────────────────╯

    # salva il .docx con le correzioni
    doc.save(out)
    print(f"💾  Documento salvato: {out.name}")

    # footnote e report restano sincroni
    client_sync = OpenAI(api_key=OPENAI_API_KEY)
    correggi_footnotes_xml(out, client_sync)
    write_markdown_report(mods, out)

def find_latest_docx(folder: Path) -> Path:
    files = list(folder.glob("*.docx"))
    if not files:
        raise RuntimeError("Nessun .docx trovato nella cartella")
    return max(files, key=lambda p: p.stat().st_mtime)

# ------------------------------------------------------------------ #
#  API pubblica per FastAPI

async def run_proofread(src: Path) -> Path:
    """
    Corregge <src>.docx e restituisce il path di uno ZIP che contiene:
      • <nome>_corretto.docx
      • <nome>_diff.md
      • <nome>_glossario.md
    """
    if not src.exists():
        raise FileNotFoundError(src)

    dst = src.with_stem(src.stem + "_corretto")
    await process_doc_async(src, dst)                      # funzione già presente

    zip_path = dst.with_suffix(".zip")
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for f in (
            dst,
            dst.with_name(dst.stem + "_diff.md"),
            dst.with_name(dst.stem + "_glossario.md"),
        ):
            if f.exists():
                zf.write(f, arcname=f.name)
    return zip_path
# ------------------------------------------------------------------ #

if __name__ == "__main__":
    # marca inizio
    start_time = time.perf_counter()

    here = Path(__file__).resolve().parent
    src = find_latest_docx(here)
    dst = src.with_stem(src.stem + "_corretto")
    print(f"📝  Correggo {src.name} → {dst.name} …")
    process_doc(src, dst)

    # tempo impiegato
    elapsed = time.perf_counter() - start_time
    print(f"✨  Fatto in {elapsed:.2f} secondi!")